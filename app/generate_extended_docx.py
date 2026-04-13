import subprocess
import sys
import os

def install(package):
    subprocess.check_call([sys.executable, "-m", "pip", "install", package])

try:
    import docx
except ImportError:
    install('python-docx')
    import docx

from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn

doc = docx.Document()

def add_rtl_p(text, is_heading=False, level=1, bold=False):
    if is_heading:
        p = doc.add_heading(text, level=level)
    else:
        p = doc.add_paragraph()
        run = p.add_run(text)
        if bold:
            run.bold = True
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

# מחולל התוכן
add_rtl_p('מסמך אפיון וארכיטקטורה - מערכת לימוד דיקדוק למקרא ועדות ישראל', True, 0)
add_rtl_p('גרסה 4 - ארכיטקטורה מורחבת', False, bold=True)
add_rtl_p('אפריל 2026', False)

# חלק א' (בקצרה, כי המיקוד הוא על ב')
add_rtl_p('חלק א׳ – תמצית פדגוגית', True, 1)
add_rtl_p('האפליקציה מתבססת על פדגוגיה חזותית יהודית: מלכים (תנועות), שווא (תמרורי עצור/זז), דגש קל (מגן), דגש חזק (מכונת צילום). כל אלו נלמדים לפי חמשת כללי אליהו בחור לשווא, ללא התערבות האקדמיה ללשון.')

# חלק ב' הארכיטקטורה המורחבת
add_rtl_p('חלק ב׳ – ארכיטקטורה טכנית מורחבת', True, 1)

add_rtl_p('1. סקירת המערכת (System Overview)', True, 2)
add_rtl_p('המערכת תיבנה כיישוום רשת מתקדם (PWA - Progressive Web App). בחירה זו מאפשרת חוויית משתמש חלקה כמו אפליקציה נייטיב (התקנה על מסך הבית), בד בבד עם יכולת עבודה באופליין (Offline-first). זה קריטי עבור משתמשים שרוצים ללמוד תוך כדי תנועה או שומרי שבת/חג שמשתמשים במכשירים ייעודיים ללא רשת זמינה באופן תמידי.\n'
          '• צד משתמש (Frontend): React.js בשילוב TypeScript לטיפוסים קפדניים.\n'
          '• ניהול מצב גלובלי (State Management): Zustand.\n'
          '• עיצוב (Styling): TailwindCSS עם תמיכה מלאה ב-RTL וערכות נושא (מצב ילדים צבעוני מול מצב מבוגרים סולידי).\n'
          '• מסד נתונים מקומי: IndexedDB באמצעות Dexie.js (לאחסון התקדמות המשתמש).')

add_rtl_p('2. מנוע הליבה: צליחת התחביר המקראי (Grammar & AST Engine)', True, 2)
add_rtl_p('המנוע אינו מחפש מחרוזות פשוטות, אלא מפרק כל פסוק או מילה ל"עץ תחביר" (AST) המייצג את המפה הווייזואלית של הדיקדוק. הניתוח מתבצע ברמת "ההקשר המשפטי" (Context-Aware), כלומר מילה מושפעת מהמילה שלפניה.')
add_rtl_p('מחסנית הניתוח (Parsing Pipeline):')
add_rtl_p('1. זיהוי בסיס (Tokenization): פירוק מילה לאות עיצור, ניקוד (Vowel) וסימנים מיוחדים (דגש, מתג, מפיק).\n'
          '2. סיווג מלכים: סיווג הניקוד ל"מלך גדול", "מלך קטן" או "חטף" לפי מסורת היעד.\n'
          '3. מפענח שוואים (Shva Resolver): מריץ את חמשת הכללים באופן רקורסיבי: בודק אות ראשונה, אות קודמת, מלך קודם, מזהה דגש מקביל (מכונת צילום) או תאומות מיד אחריו.\n'
          '4. מפענח דגשים (Dagesh Resolver): מבדיל בין דגש קל (מגן) מדגש חזק. בודק את המילה הקודמת (Previous Word) - האם היא הסתיימה בתנועה גדולה? האם מחוברת במקף?')
add_rtl_p('ממשק (Interface) לדוגמה:', bold=True)
add_rtl_p('export interface HebrewToken {\n'
          '  consonant: string; // "ב" \n'
          '  vowelType?: "bigKing" | "smallKing" | "halfVowel"; // "פתח" = smallKing\n'
          '  shvaState?: "stop" | "run" | "none"; // בממשק: 🛑 או 🏃‍♂️\n'
          '  dageshType?: "shield" | "photocopier" | "none"; // 🛡️ או 🖨️\n'
          '}')

add_rtl_p('3. מסדי הנתונים ומודל התוכן (Data Models)', True, 2)
add_rtl_p('התוכן נטען מקובצי JSON יעילים שיקודדו אל תוך קובצי האפליקציה (כדי להימנע מתלות בשרת).')
add_rtl_p('א. סכמת שיעור (Lesson Schema):')
add_rtl_p('כל שיעור מורכב משלבי הסבר (עם דימויים), ושורת תרגילים אינטראקטיביים. למשל תרגיל "זיהוי תמרורים", בו המנוע מקבל מחרוזת מנוקדת מקראית אמיתית, והלומד צריך להרכיב את הסטטוס לקריאה ע"י גרירת תמרורים. ה-Grammar Engine יבדוק את תשובת התלמיד מול התשובה האלגוריתמית.')
add_rtl_p('ב. סכמת מסורות (Traditions Config):')
add_rtl_p('{\"id\": \"teiman\", \"features\": {\"dageshGimelRafe\": true, \"kamatzAsCholam\": true, \"shvaNaSound\": \"variable\"}} - הגדרות אלו משפיעות ישירות על האודיו המופק ועל כללי האזהרות.')

add_rtl_p('4. מנוע המשחוק והמשוב הדינמי (Gamification & Feedback)', True, 2)
add_rtl_p('המערכת מעניקה פלפול מיידי על תשובה שגויה. במקום לכתוב "שגיאה", ה-Engine מקשר לדימוי.\n'
          '• תרחיש: תלמיד שם דגש חזק באות ע\'. המנוע מזהה (`isGuttural: true`) ופולט משוב: "שכחת? האות ע\' היא מורדת! מכונת הצילום לא מצליחה לעבוד עליה!".\n'
          '• תרחיש: תלמיד סימן "שווא נח" (תמרור עצור 🛑) אחרי מלך גדול. המשוב: "שים לב! לפניך עומד מלך גדול וחזק (קמץ). הוא משחרר את השווא לחופשי (לנוע) 🏃‍♂️!".')

add_rtl_p('5. מנוע השמע (Audio Architecture)', True, 2)
add_rtl_p('מכיוון שהפרויקט מכוון לדיוק יהודי עתיק, לא ניתן להסתמך על רכיבי ה-Text-to-Speech הרגילים של הדפדפן (שקוראים בעברית מודרנית ומשבשים שוואים). \n'
          'פתרון: בניית Audio Engine שמייצר Audio Sprites (חיתוכים של צלילי הברות) ומחבר אותם בזמן אמת (Concatenative Synthesis) בהתאם למודל המנוע המדוקדק. תהיה ערכת סאונד נפרדת לספרדי (ירושלמי), אשכנזי ותימני.')

add_rtl_p('6. פריסה ואבטחת נתונים (CI/CD & Security)', True, 2)
add_rtl_p('• אבטחה: העדר מסד בינלאומי (שרת) מבטיח 100% הגנה על פרטיות הילד. הנתונים (מטבעות, התקדמות) נשארים על הדפדפן.\n'
          '• פריסה (Deployment): שימוש ב-GitHub Actions לבדיקות אוטומטיות ל-Grammar Engine מול אלפי תרחישי מקרא (פסוקים מתויגים). פריסה ל-Vercel או Cloudflare Pages כנכסים סטטיים (Files).')

output_path = r'd:\Oren_In_D\VS_CODE\ניקוד_עברית\אפיון_ארכיטקטורה_דיקדוק_מורחב.docx'
doc.save(output_path)
print(f"Generated extended docx at: {output_path}")
